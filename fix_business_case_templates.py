#!/usr/bin/env python3
"""
Fix Business Case Templates
Regenerates all 31 mismatched Business Case templates with correct content
"""

import os
import sys
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from app import app, db
from models import Template

# List of mismatched Business Case template IDs from audit
MISMATCHED_IDS = [
    1, 34, 65, 95, 127, 158, 188, 222, 253, 289,
    319, 350, 381, 412, 444, 476, 506, 539, 570, 600,
    634, 666, 699, 733, 762, 799, 829, 862, 894, 926
]

def get_industry_from_filename(filename):
    """Extract industry name from filename"""
    # Remove extension and _Business_Case suffix
    name = filename.replace('.docx', '').replace('_Business_Case', '')
    # Replace underscores with spaces
    name = name.replace('_', ' ')
    return name

def generate_business_case_content(industry):
    """Generate Business Case document content for an industry"""
    
    content = f"""BUSINESS CASE
{industry.upper()} PROJECT

Strategic Business Justification and Investment Analysis

Document Type: Business Case
Industry: {industry}
Date: October 20, 2025
Version: 1.0
PMI Standard: 2025 Compliance
Classification: Confidential - Executive Level

EXECUTIVE SUMMARY

This business case presents a comprehensive analysis and justification for a strategic {industry} initiative. The proposed project addresses critical business challenges and opportunities while delivering measurable value to the organization.

Project Overview:
• Industry Focus: {industry}
• Strategic Alignment: Supports organizational objectives and competitive positioning
• Investment Type: Strategic initiative with measurable ROI
• Timeline: 12-24 months implementation
• Risk Level: Medium (manageable with proper controls)

Key Benefits:
• Operational efficiency improvements
• Cost reduction and optimization
• Revenue growth opportunities
• Competitive advantage
• Risk mitigation

Financial Summary:
• Total Investment Required: $2,500,000 - $5,000,000
• Expected ROI: 200-300% over 3 years
• Payback Period: 18-24 months
• NPV (Net Present Value): Positive
• IRR (Internal Rate of Return): 25-35%

This initiative aligns with organizational strategic objectives and industry best practices, positioning the organization for sustainable growth and competitive advantage.

BUSINESS PROBLEM / OPPORTUNITY

Current State Analysis:
The organization faces several challenges and opportunities in the {industry} domain that require strategic intervention:

Challenges:
• Operational inefficiencies impacting productivity
• Increasing competitive pressure
• Changing market dynamics and customer expectations
• Technology gaps and legacy system limitations
• Regulatory and compliance requirements

Opportunities:
• Market expansion potential
• Process optimization and automation
• Enhanced customer experience and satisfaction
• Data-driven decision making capabilities
• Innovation and competitive differentiation

Impact of Inaction:
Failure to address these challenges will result in:
• Loss of market share to competitors
• Declining operational efficiency
• Increased costs and reduced profitability
• Customer dissatisfaction and attrition
• Regulatory compliance risks

PROPOSED SOLUTION

Solution Overview:
The proposed solution involves a comprehensive {industry} initiative that addresses the identified challenges through:

Key Components:
1. Strategic Planning and Assessment
   • Current state analysis and gap assessment
   • Future state design and roadmap development
   • Stakeholder alignment and buy-in

2. Implementation and Execution
   • Phased rollout approach
   • Change management and training
   • Quality assurance and testing

3. Optimization and Continuous Improvement
   • Performance monitoring and measurement
   • Ongoing optimization and refinement
   • Knowledge transfer and sustainability

Strategic Alignment:
This solution directly supports:
• Corporate strategic objectives
• Industry best practices and standards
• Regulatory compliance requirements
• Customer satisfaction goals
• Financial performance targets

COST-BENEFIT ANALYSIS

Investment Requirements:

Initial Investment:
• Technology and Infrastructure: $1,500,000
• Implementation Services: $800,000
• Training and Change Management: $400,000
• Contingency (15%): $405,000
• Total Initial Investment: $3,105,000

Ongoing Costs (Annual):
• Operations and Maintenance: $250,000
• Support and Updates: $150,000
• Training and Development: $100,000
• Total Annual Costs: $500,000

Expected Benefits:

Quantifiable Benefits (Annual):
• Cost Savings: $1,200,000
• Revenue Increase: $800,000
• Efficiency Gains: $600,000
• Total Annual Benefits: $2,600,000

Intangible Benefits:
• Improved customer satisfaction and loyalty
• Enhanced employee productivity and morale
• Stronger competitive positioning
• Better risk management and compliance
• Increased organizational agility

Financial Metrics:

ROI Analysis (3-Year Period):
• Total Investment: $4,605,000
• Total Benefits: $7,800,000
• Net Benefit: $3,195,000
• ROI: 69% (3-year cumulative)
• Annual ROI: 23%

Payback Period: 21 months

Net Present Value (NPV):
• Discount Rate: 10%
• NPV: $2,450,000 (positive)

Internal Rate of Return (IRR): 28%

RISK ASSESSMENT

Risk Category: Medium

Key Risks and Mitigation Strategies:

1. Implementation Risk (Medium)
   • Risk: Project delays or scope creep
   • Mitigation: Robust project management, clear governance, phased approach
   • Contingency: Additional resources and timeline buffers

2. Technology Risk (Medium)
   • Risk: Technical challenges or integration issues
   • Mitigation: Thorough technical assessment, proof of concept, vendor support
   • Contingency: Alternative solutions and rollback procedures

3. Change Management Risk (High)
   • Risk: User resistance or adoption challenges
   • Mitigation: Comprehensive change management program, training, communication
   • Contingency: Extended support period and additional training

4. Financial Risk (Low)
   • Risk: Cost overruns or lower than expected benefits
   • Mitigation: Detailed budgeting, regular monitoring, benefit tracking
   • Contingency: Cost controls and benefit realization management

5. Vendor/Partner Risk (Low)
   • Risk: Vendor performance or relationship issues
   • Mitigation: Thorough vendor selection, clear contracts, performance monitoring
   • Contingency: Alternative vendor options and exit strategies

Overall Risk Rating: MEDIUM (Acceptable with proper mitigation)

IMPLEMENTATION APPROACH

Phased Implementation Strategy:

Phase 1: Planning and Design (Months 1-3)
• Detailed requirements gathering
• Solution design and architecture
• Vendor selection and contracting
• Project team formation
• Stakeholder engagement

Phase 2: Development and Testing (Months 4-9)
• Solution development and configuration
• Integration and data migration
• Testing and quality assurance
• Training material development
• Pilot preparation

Phase 3: Pilot and Refinement (Months 10-12)
• Pilot deployment
• User feedback and refinement
• Performance monitoring
• Issue resolution
• Go-live preparation

Phase 4: Full Deployment (Months 13-18)
• Phased rollout to all users
• Comprehensive training delivery
• Change management execution
• Support and stabilization
• Performance measurement

Phase 5: Optimization (Months 19-24)
• Continuous improvement
• Benefit realization tracking
• Lessons learned documentation
• Knowledge transfer
• Sustainability planning

RECOMMENDATION

Based on the comprehensive analysis presented in this business case, we recommend proceeding with this {industry} initiative for the following reasons:

1. Strong Financial Justification
   • Positive ROI of 69% over 3 years
   • Payback period of 21 months
   • Positive NPV and strong IRR

2. Strategic Alignment
   • Directly supports organizational objectives
   • Addresses critical business challenges
   • Positions organization for competitive advantage

3. Manageable Risk Profile
   • Medium risk level with effective mitigation strategies
   • Proven implementation approach
   • Strong governance and oversight

4. Measurable Benefits
   • Clear quantifiable benefits
   • Significant intangible value
   • Comprehensive benefit realization plan

5. Stakeholder Support
   • Executive sponsorship
   • Business unit alignment
   • User community engagement

NEXT STEPS

Upon approval of this business case, the following immediate actions are recommended:

1. Secure Executive Approval and Funding (Week 1-2)
   • Present to executive committee
   • Obtain budget allocation
   • Formalize project charter

2. Establish Project Governance (Week 2-3)
   • Form steering committee
   • Appoint project sponsor
   • Define decision-making framework

3. Initiate Project Planning (Week 3-6)
   • Assemble project team
   • Develop detailed project plan
   • Begin vendor selection process

4. Commence Phase 1 Activities (Week 7+)
   • Launch requirements gathering
   • Initiate stakeholder engagement
   • Begin solution design

APPROVAL

This business case requires approval from:

Executive Sponsor: _____________________________ Date: __________

CFO: _____________________________ Date: __________

CIO/CTO: _____________________________ Date: __________

Business Unit Leader: _____________________________ Date: __________

---

Document prepared by: Project Management Office
Date: October 20, 2025
Version: 1.0
Classification: Confidential - Executive Level
"""
    
    return content

def create_business_case_docx(industry, output_path):
    """Create a Business Case DOCX file"""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    content = generate_business_case_content(industry)
    
    # Split content into sections
    sections_text = content.split('\n\n')
    
    for section in sections_text:
        if not section.strip():
            continue
            
        lines = section.split('\n')
        
        # Check if this is a heading (all caps, short line)
        if len(lines) == 1 and lines[0].isupper() and len(lines[0]) < 80:
            # Add as heading
            heading = doc.add_heading(lines[0], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Add as paragraph(s)
            for line in lines:
                if line.strip():
                    if line.startswith('•'):
                        # Bullet point
                        p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                    elif ':' in line and len(line.split(':')[0]) < 50:
                        # Key-value pair
                        p = doc.add_paragraph()
                        parts = line.split(':', 1)
                        run = p.add_run(parts[0] + ':')
                        run.bold = True
                        if len(parts) > 1:
                            p.add_run(parts[1])
                    else:
                        # Regular paragraph
                        doc.add_paragraph(line)
    
    doc.save(output_path)
    print(f"✅ Created: {output_path}")

def main():
    print("=" * 80)
    print("FIXING BUSINESS CASE TEMPLATES")
    print("=" * 80)
    print()
    
    templates_dir = Path(__file__).parent / 'static' / 'templates'
    
    with app.app_context():
        for template_id in MISMATCHED_IDS:
            template = Template.query.get(template_id)
            
            if not template:
                print(f"❌ Template ID {template_id} not found")
                continue
            
            print(f"Fixing ID {template.id}: {template.name}")
            print(f"  File: {template.file_path}")
            
            # Extract industry from filename
            industry = get_industry_from_filename(template.file_path)
            print(f"  Industry: {industry}")
            
            # Generate new Business Case document
            output_path = templates_dir / template.file_path
            
            try:
                create_business_case_docx(industry, output_path)
            except Exception as e:
                print(f"❌ Error creating {template.file_path}: {e}")
                continue
    
    print()
    print("=" * 80)
    print("COMPLETE - All Business Case templates regenerated")
    print("=" * 80)

if __name__ == '__main__':
    main()

