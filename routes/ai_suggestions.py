"""
AI Suggestions Routes
Handles AI-powered template suggestions - optimized for fast ChatGPT-like responses
"""

from flask import Blueprint, render_template, request, jsonify
from flask_login import login_required, current_user
import logging
import os
from openai import OpenAI
from datetime import datetime

logger = logging.getLogger(__name__)

ai_suggestions_bp = Blueprint('ai_suggestions', __name__)

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

@ai_suggestions_bp.route('/')
def index():
    """AI Suggestions page - public view, suggestions require login"""
    from flask_login import current_user
    from utils.subscription_security import get_usage_stats
    
    usage_stats = get_usage_stats(current_user) if current_user.is_authenticated else None
    
    return render_template('ai_suggestions.html', usage_stats=usage_stats)

@ai_suggestions_bp.route('/get', methods=['POST'])
@login_required
def get_suggestions():
    """Get AI-powered template suggestions - fast response optimized"""
    from database import db
    from models import AISuggestionHistory
    
    try:
        # Get request data
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        # Extract user inputs
        template_type = data.get('template_type', '').strip()
        section = data.get('section', '').strip()
        question = data.get('question', '').strip()
        context = data.get('context', '').strip()
        
        # Validate inputs
        if not template_type:
            return jsonify({'error': 'Template type is required'}), 400
        
        # Build prompt based on request type
        if section:
            # Quick suggestion format
            user_prompt = f"For a {template_type} project, provide specific suggestions for: {section}"
            if context:
                user_prompt += f"\n\nAdditional context: {context}"
        elif question:
            # Custom question format
            user_prompt = f"Template Type: {template_type}\n\nQuestion: {question}"
            if context:
                user_prompt += f"\n\nContext: {context}"
        else:
            return jsonify({'error': 'Either section or question is required'}), 400
        
        # Call OpenAI API with optimized settings for speed
        response = client.chat.completions.create(
            model="gpt-4o-mini",  # Fast, efficient model
            messages=[
                {
                    "role": "system", 
                    "content": "You are an expert project management consultant. Provide concise, actionable advice based on PMBOK standards and industry best practices. Be direct and specific."
                },
                {
                    "role": "user", 
                    "content": user_prompt
                }
            ],
            temperature=0.7,
            max_tokens=800,  # Reduced for faster response
            stream=False  # No streaming for simpler implementation
        )
        
        suggestions_text = response.choices[0].message.content
        
        # Save suggestion history (async-friendly, minimal blocking)
        try:
            history = AISuggestionHistory(
                user_id=current_user.id,
                project_description=question or f"{template_type} - {section}",
                industry=template_type,
                project_phase='',
                team_size=context or '',
                suggestions=suggestions_text,
                created_at=datetime.utcnow()
            )
            db.session.add(history)
            db.session.commit()
            history_id = history.id
        except Exception as db_error:
            logger.error(f"Database error saving history: {str(db_error)}")
            # Don't fail the request if history save fails
            history_id = None
        
        # Return success immediately
        return jsonify({
            'success': True,
            'suggestion': suggestions_text,
            'history_id': history_id
        })
        
    except Exception as e:
        logger.error(f"AI suggestions error: {str(e)}")
        
        # Handle specific errors
        error_message = str(e)
        if 'rate_limit' in error_message.lower() or 'quota' in error_message.lower():
            return jsonify({'error': 'AI service is currently busy. Please try again in a moment.'}), 429
        elif 'authentication' in error_message.lower() or 'api_key' in error_message.lower():
            return jsonify({'error': 'AI service configuration error. Please contact support.'}), 500
        else:
            return jsonify({'error': 'An error occurred while generating suggestions. Please try again.'}), 500

@ai_suggestions_bp.route('/history')
@login_required
def history():
    """View suggestion history"""
    from models import AISuggestionHistory
    
    page = request.args.get('page', 1, type=int)
    per_page = 20
    
    history_query = AISuggestionHistory.query.filter_by(
        user_id=current_user.id
    ).order_by(AISuggestionHistory.created_at.desc())
    
    pagination = history_query.paginate(page=page, per_page=per_page, error_out=False)
    
    return render_template('ai/suggestions_history.html', 
                         history=pagination.items,
                         pagination=pagination)

@ai_suggestions_bp.route('/export/<int:history_id>')
@login_required
def export(history_id):
    """Export suggestions as Word document"""
    from models import AISuggestionHistory
    from docx import Document
    import io
    from flask import send_file
    
    history = AISuggestionHistory.query.get_or_404(history_id)
    
    # Verify ownership
    if history.user_id != current_user.id:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        # Create Word document
        doc = Document()
        doc.add_heading('AI Template Suggestions', 0)
        
        doc.add_heading('Project Details', level=1)
        doc.add_paragraph(f"Description: {history.project_description}")
        doc.add_paragraph(f"Industry: {history.industry or 'Not specified'}")
        doc.add_paragraph(f"Team Size/Context: {history.team_size or 'Not specified'}")
        doc.add_paragraph(f"Generated: {history.created_at.strftime('%Y-%m-%d %H:%M')}")
        
        doc.add_heading('AI Suggestions', level=1)
        doc.add_paragraph(history.suggestions)
        
        # Save to BytesIO
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        filename = f"ai_suggestions_{history.id}.docx"
        
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logger.error(f"Export error: {str(e)}")
        return jsonify({'error': 'Failed to export suggestions'}), 500

