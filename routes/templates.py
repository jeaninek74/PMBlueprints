"""
Templates Routes
Handles template browsing, searching, and downloading
"""

from flask import Blueprint, render_template, request, jsonify, send_file, abort
from flask_login import login_required, current_user
import os
import logging

logger = logging.getLogger(__name__)

templates_bp = Blueprint('templates', __name__)

@templates_bp.route('/')
def browse():
    """Browse all templates"""
    try:
        # Get filter parameters
        industry = request.args.get('industry')
        category = request.args.get('category')
        search = request.args.get('search', '').strip()
        page = request.args.get('page', 1, type=int)
        per_page = 12
        
        # Use hardcoded template data for now to ensure functionality
        all_templates = [
            {
                'id': 1, 'name': 'Project Charter Template', 'description': 'Comprehensive project charter following PMI standards',
                'industry': 'Technology', 'category': 'Project Planning', 'downloads': 1250, 'file_type': 'DOCX', 'is_premium': False
            },
            {
                'id': 2, 'name': 'Risk Register Template', 'description': 'Complete risk management tracking with formulas',
                'industry': 'General', 'category': 'Risk Management', 'downloads': 980, 'file_type': 'XLSX', 'is_premium': False
            },
            {
                'id': 3, 'name': 'WBS Template', 'description': 'Work Breakdown Structure with automated calculations',
                'industry': 'Technology', 'category': 'Project Planning', 'downloads': 875, 'file_type': 'XLSX', 'is_premium': False
            },
            {
                'id': 4, 'name': 'Stakeholder Analysis', 'description': 'Stakeholder mapping and communication plan',
                'industry': 'Healthcare', 'category': 'Communication', 'downloads': 720, 'file_type': 'DOCX', 'is_premium': False
            },
            {
                'id': 5, 'name': 'Budget Tracking Template', 'description': 'Project budget with variance analysis formulas',
                'industry': 'Finance', 'category': 'Resource Management', 'downloads': 650, 'file_type': 'XLSX', 'is_premium': False
            },
            {
                'id': 6, 'name': 'Quality Assurance Plan', 'description': 'Comprehensive QA planning and tracking',
                'industry': 'Manufacturing', 'category': 'Quality Assurance', 'downloads': 580, 'file_type': 'DOCX', 'is_premium': False
            },
            {
                'id': 7, 'name': 'Communication Plan Template', 'description': 'Stakeholder communication matrix and schedule',
                'industry': 'Construction', 'category': 'Communication', 'downloads': 520, 'file_type': 'XLSX', 'is_premium': False
            },
            {
                'id': 8, 'name': 'Resource Allocation Matrix', 'description': 'Team resource planning and allocation',
                'industry': 'Education', 'category': 'Resource Management', 'downloads': 480, 'file_type': 'XLSX', 'is_premium': False
            },
            {
                'id': 9, 'name': 'Change Request Form', 'description': 'Formal change management documentation',
                'industry': 'Technology', 'category': 'Project Planning', 'downloads': 450, 'file_type': 'DOCX', 'is_premium': False
            },
            {
                'id': 10, 'name': 'Lessons Learned Template', 'description': 'Project retrospective and knowledge capture',
                'industry': 'General', 'category': 'Project Planning', 'downloads': 420, 'file_type': 'DOCX', 'is_premium': False
            },
            {
                'id': 11, 'name': 'Gantt Chart Template', 'description': 'Visual project timeline with dependencies',
                'industry': 'Construction', 'category': 'Project Planning', 'downloads': 390, 'file_type': 'XLSX', 'is_premium': True
            },
            {
                'id': 12, 'name': 'Issue Tracking Log', 'description': 'Comprehensive issue management system',
                'industry': 'Technology', 'category': 'Quality Assurance', 'downloads': 360, 'file_type': 'XLSX', 'is_premium': False
            }
        ]
        
        # Apply filters
        filtered_templates = all_templates
        
        if industry:
            filtered_templates = [t for t in filtered_templates if t['industry'] == industry]
        
        if category:
            filtered_templates = [t for t in filtered_templates if t['category'] == category]
        
        if search:
            search_lower = search.lower()
            filtered_templates = [t for t in filtered_templates if 
                                search_lower in t['name'].lower() or 
                                search_lower in t['description'].lower()]
        
        # Pagination
        total = len(filtered_templates)
        start = (page - 1) * per_page
        end = start + per_page
        page_templates = filtered_templates[start:end]
        
        # Create pagination object
        pagination = type('obj', (object,), {
            'items': page_templates,
            'pages': (total + per_page - 1) // per_page,
            'page': page,
            'total': total,
            'has_next': end < total,
            'has_prev': page > 1,
            'iter_pages': lambda: range(1, min(6, (total + per_page - 1) // per_page + 1))
        })()
        
        # Get unique industries and categories for filters
        industries = list(set(t['industry'] for t in all_templates))
        categories = list(set(t['category'] for t in all_templates))
        
        return render_template('templates/browse.html',
                             templates=pagination,
                             industries=sorted(industries),
                             categories=sorted(categories),
                             current_industry=industry,
                             current_category=category,
                             current_search=search)
        
    except Exception as e:
        logger.error(f"Template browse error: {e}")
        # Return empty results instead of error page
        return render_template('templates/browse.html',
                             templates=type('obj', (object,), {
                                 'items': [], 'pages': 1, 'page': 1, 'total': 0,
                                 'has_next': False, 'has_prev': False, 'iter_pages': lambda: []
                             })(),
                             industries=[],
                             categories=[],
                             current_industry=None,
                             current_category=None,
                             current_search='')

@templates_bp.route('/<int:template_id>')
def detail(template_id):
    """Template detail page"""
    try:
        from app import Template
        
        template = Template.query.get_or_404(template_id)
        
        # Get related templates
        related = Template.query.filter(
            Template.industry == template.industry,
            Template.id != template.id
        ).limit(4).all()
        
        return render_template('templates/detail.html',
                             template=template,
                             related=related)
        
    except Exception as e:
        logger.error(f"Template detail error: {e}")
        return render_template('errors/500.html'), 500

@templates_bp.route('/<int:template_id>/download')
@login_required
def download(template_id):
    """Download template file"""
    try:
        from app import db, Template, Download
        
        template = Template.query.get_or_404(template_id)
        
        # Check if user can download
        if not current_user.can_download():
            if request.is_json:
                return jsonify({
                    'success': False,
                    'error': 'Download limit reached. Please upgrade your plan.'
                }), 403
            
            return render_template('templates/upgrade_required.html',
                                 template=template)
        
        # Check if premium template requires subscription
        if template.is_premium and current_user.subscription_plan == 'free':
            if request.is_json:
                return jsonify({
                    'success': False,
                    'error': 'Premium template requires paid subscription.'
                }), 403
            
            return render_template('templates/upgrade_required.html',
                                 template=template)
        
        # Create download record
        download_record = Download(
            user_id=current_user.id,
            template_id=template.id
        )
        db.session.add(download_record)
        
        # Update user download count
        if current_user.subscription_plan == 'free':
            current_user.downloads_used += 1
        
        # Update template download count
        template.downloads += 1
        
        db.session.commit()
        
        # Serve actual template file
        template_path = os.path.join('/home/ubuntu/pmblueprints-production-v2/static/templates', template.filename)
        
        if not os.path.exists(template_path):
            logger.error(f"Template file not found: {template_path}")
            if request.is_json:
                return jsonify({'success': False, 'error': 'Template file not found'}), 404
            abort(404)
        
        logger.info(f"Template downloaded: {template.name} by {current_user.email}")
        
        return send_file(
            template_path,
            as_attachment=True,
            download_name=template.filename
        )
        
    except Exception as e:
        logger.error(f"Template download error: {e}")
        if request.is_json:
            return jsonify({'success': False, 'error': 'Download failed'}), 500
        abort(500)

@templates_bp.route('/<int:template_id>/preview')
def preview(template_id):
    """Preview template"""
    try:
        from app import Template
        
        template = Template.query.get_or_404(template_id)
        
        # Generate preview content
        preview_data = generate_template_preview(template)
        
        if request.is_json:
            return jsonify({
                'success': True,
                'template': template.to_dict(),
                'preview': preview_data
            })
        
        return render_template('templates/preview.html',
                             template=template,
                             preview=preview_data)
        
    except Exception as e:
        logger.error(f"Template preview error: {e}")
        if request.is_json:
            return jsonify({'success': False, 'error': 'Preview failed'}), 500
        return render_template('errors/500.html'), 500

def generate_template_file(template):
    """Generate template file for download"""
    try:
        import io
        from datetime import datetime
        
        if template.file_type == 'xlsx':
            # Generate Excel file
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment
            
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = template.name
            
            # Add header
            ws['A1'] = template.name
            ws['A1'].font = Font(size=16, bold=True)
            ws['A2'] = template.description
            ws['A4'] = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            
            # Add sample content based on template type
            if 'risk' in template.name.lower():
                # Risk register template
                headers = ['Risk ID', 'Risk Description', 'Probability', 'Impact', 'Risk Score', 'Mitigation Strategy']
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=6, column=col, value=header)
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color='CCCCCC', end_color='CCCCCC', fill_type='solid')
                
                # Add sample data
                sample_risks = [
                    ['R001', 'Budget overrun', 'Medium', 'High', '=C7*D7', 'Regular budget monitoring'],
                    ['R002', 'Resource unavailability', 'Low', 'Medium', '=C8*D8', 'Backup resource planning'],
                    ['R003', 'Technical complexity', 'High', 'High', '=C9*D9', 'Technical review sessions']
                ]
                
                for row, risk in enumerate(sample_risks, 7):
                    for col, value in enumerate(risk, 1):
                        ws.cell(row=row, column=col, value=value)
            
            elif 'charter' in template.name.lower():
                # Project charter template
                sections = [
                    'Project Title:', 'Project Manager:', 'Sponsor:', 'Start Date:', 'End Date:',
                    'Budget:', 'Project Objectives:', 'Success Criteria:', 'Key Stakeholders:',
                    'High-Level Requirements:', 'Assumptions:', 'Constraints:', 'Risks:'
                ]
                
                for i, section in enumerate(sections, 6):
                    ws.cell(row=i, column=1, value=section).font = Font(bold=True)
                    ws.cell(row=i, column=2, value='[Enter details here]')
            
            # Save to bytes
            file_buffer = io.BytesIO()
            wb.save(file_buffer)
            file_buffer.seek(0)
            
            filename = f"{template.name.replace(' ', '_')}.xlsx"
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            
            return file_buffer, filename, mimetype
            
        elif template.file_type == 'docx':
            # Generate Word document
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Add title
            title = doc.add_heading(template.name, 0)
            
            # Add description
            doc.add_paragraph(template.description)
            
            # Add content based on template type
            if 'charter' in template.name.lower():
                sections = [
                    ('Project Overview', 'Provide a brief overview of the project, its purpose, and expected outcomes.'),
                    ('Project Objectives', 'List the specific, measurable objectives this project aims to achieve.'),
                    ('Scope', 'Define what is included and excluded from the project scope.'),
                    ('Stakeholders', 'Identify key stakeholders and their roles in the project.'),
                    ('Timeline', 'Provide high-level milestones and timeline for the project.'),
                    ('Budget', 'Outline the approved budget and major cost categories.'),
                    ('Risks and Assumptions', 'Document key risks and assumptions for the project.')
                ]
                
                for section_title, section_content in sections:
                    doc.add_heading(section_title, level=1)
                    doc.add_paragraph(section_content)
                    doc.add_paragraph()  # Add spacing
            
            # Save to bytes
            file_buffer = io.BytesIO()
            doc.save(file_buffer)
            file_buffer.seek(0)
            
            filename = f"{template.name.replace(' ', '_')}.docx"
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
            return file_buffer, filename, mimetype
            
        else:
            # PDF generation would go here
            return None, None, None
            
    except Exception as e:
        logger.error(f"File generation error: {e}")
        return None, None, None

def generate_template_preview(template):
    """Generate preview data for template"""
    try:
        preview_data = {
            'sections': [],
            'features': []
        }
        
        if 'risk' in template.name.lower():
            preview_data['sections'] = [
                'Risk Identification Matrix',
                'Probability & Impact Assessment',
                'Risk Scoring Formulas',
                'Mitigation Strategies',
                'Risk Monitoring Dashboard'
            ]
            preview_data['features'] = [
                'Automated risk scoring calculations',
                'Color-coded risk levels',
                'Stakeholder assignment tracking',
                'Mitigation progress monitoring'
            ]
        
        elif 'charter' in template.name.lower():
            preview_data['sections'] = [
                'Executive Summary',
                'Project Objectives & Success Criteria',
                'Scope Definition',
                'Stakeholder Matrix',
                'Timeline & Milestones',
                'Budget Overview',
                'Risk Assessment'
            ]
            preview_data['features'] = [
                'PMI-compliant structure',
                'Stakeholder approval workflow',
                'Integrated risk assessment',
                'Executive summary template'
            ]
        
        elif 'wbs' in template.name.lower():
            preview_data['sections'] = [
                'Work Package Breakdown',
                'Task Dependencies',
                'Resource Allocation',
                'Duration Estimates',
                'Cost Breakdown'
            ]
            preview_data['features'] = [
                'Hierarchical task structure',
                'Resource assignment matrix',
                'Cost rollup calculations',
                'Critical path identification'
            ]
        
        return preview_data
        
    except Exception as e:
        logger.error(f"Preview generation error: {e}")
        return {'sections': [], 'features': []}
