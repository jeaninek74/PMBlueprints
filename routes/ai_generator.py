"""
AI Generator Routes
Handles AI-powered document generation with tier-based access control
"""

from flask import Blueprint, render_template, request, jsonify, send_file, flash, redirect, url_for
from flask_login import login_required, current_user
import logging
import os
import openai
from datetime import datetime
from docx import Document
from openpyxl import Workbook
import io

logger = logging.getLogger(__name__)

ai_generator_bp = Blueprint('ai_generator', __name__, url_prefix='/ai/generator')

# Initialize OpenAI
openai.api_key = os.getenv('OPENAI_API_KEY')

@ai_generator_bp.route('/')
@login_required
def index():
    """AI Generator page"""
    from utils.subscription_security import get_usage_stats
    
    usage_stats = get_usage_stats(current_user)
    
    return render_template('ai/generator.html', usage_stats=usage_stats)

@ai_generator_bp.route('/generate', methods=['POST'])
@login_required
def generate():
    """Generate AI document"""
    from utils.subscription_security import check_usage_limit
    from database import db
    from models import AIGeneratorHistory
    
    try:
        # Check usage quota
        can_generate, remaining, limit = check_usage_limit(current_user, 'ai_generations')
        
        if not can_generate:
            return jsonify({
                'error': f'You have reached your AI generation limit ({limit} per month). Please upgrade your subscription.',
                'upgrade_required': True,
                'remaining': 0,
                'limit': limit
            }), 403
        
        # Get request data
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        project_name = data.get('project_name', '')
        project_type = data.get('project_type', '')
        industry = data.get('industry', '')
        methodology = data.get('methodology', 'Agile')
        document_type = data.get('document_type', 'Project Charter')
        additional_details = data.get('additional_details', '')
        
        if not project_name or not project_type or not industry:
            return jsonify({'error': 'Missing required fields'}), 400
        
        # Generate document using OpenAI
        prompt = f"""
Generate a professional {document_type} for a {methodology} project.

Project Details:
- Project Name: {project_name}
- Project Type: {project_type}
- Industry: {industry}
- Methodology: {methodology}
- Additional Details: {additional_details}

Please create a comprehensive {document_type} following PMI PMBOK standards.
Include all relevant sections, proper formatting, and industry-specific considerations.
Make it professional and ready to use.
"""
        
        # Call OpenAI API
        response = openai.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert project management consultant specializing in creating professional PM documents following PMI standards."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=3000
        )
        
        generated_content = response.choices[0].message.content
        
        # Create document file based on type
        file_format = data.get('file_format', 'docx')
        file_data, filename = create_document_file(
            generated_content,
            document_type,
            project_name,
            file_format
        )
        
        # Track usage
        current_user.ai_generations_this_month += 1
        
        # Save generation history
        history = AIGeneratorHistory(
            user_id=current_user.id,
            project_name=project_name,
            project_type=project_type,
            industry=industry,
            methodology=methodology,
            document_type=document_type,
            file_format=file_format,
            generated_content=generated_content[:5000],  # Store first 5000 chars
            file_path=filename,
            created_at=datetime.utcnow()
        )
        db.session.add(history)
        db.session.commit()
        
        # Return success with download link
        return jsonify({
            'success': True,
            'message': 'Document generated successfully',
            'history_id': history.id,
            'download_url': url_for('ai_generator.download', history_id=history.id, _external=True),
            'remaining': remaining - 1,
            'limit': limit,
            'preview': generated_content[:500] + '...'
        })
        
    except openai.error.RateLimitError:
        logger.error("OpenAI rate limit exceeded")
        return jsonify({'error': 'AI service is currently busy. Please try again in a moment.'}), 429
    
    except openai.error.AuthenticationError:
        logger.error("OpenAI authentication error")
        return jsonify({'error': 'AI service configuration error. Please contact support.'}), 500
    
    except Exception as e:
        logger.error(f"AI generation error: {str(e)}")
        db.session.rollback()
        return jsonify({'error': f'An error occurred: {str(e)}'}), 500

@ai_generator_bp.route('/download/<int:history_id>')
@login_required
def download(history_id):
    """Download generated document"""
    from models import AIGeneratorHistory
    
    history = AIGeneratorHistory.query.get_or_404(history_id)
    
    # Verify ownership
    if history.user_id != current_user.id:
        flash('Unauthorized access', 'error')
        return redirect(url_for('ai_generator.index'))
    
    try:
        # Recreate the document file
        file_data, filename = create_document_file(
            history.generated_content,
            history.document_type,
            history.project_name,
            history.file_format
        )
        
        return send_file(
            file_data,
            as_attachment=True,
            download_name=filename,
            mimetype=get_mimetype(history.file_format)
        )
        
    except Exception as e:
        logger.error(f"Download error: {str(e)}")
        flash('Error downloading file', 'error')
        return redirect(url_for('ai_generator.index'))

@ai_generator_bp.route('/history')
@login_required
def history():
    """View generation history"""
    from models import AIGeneratorHistory
    
    page = request.args.get('page', 1, type=int)
    per_page = 20
    
    history_query = AIGeneratorHistory.query.filter_by(
        user_id=current_user.id
    ).order_by(AIGeneratorHistory.created_at.desc())
    
    pagination = history_query.paginate(page=page, per_page=per_page, error_out=False)
    
    return render_template('ai/history.html', 
                         history=pagination.items,
                         pagination=pagination)

def create_document_file(content, document_type, project_name, file_format):
    """Create document file from generated content"""
    
    # Sanitize filename
    safe_project_name = "".join(c for c in project_name if c.isalnum() or c in (' ', '-', '_')).strip()
    safe_doc_type = "".join(c for c in document_type if c.isalnum() or c in (' ', '-', '_')).strip()
    
    if file_format == 'docx':
        # Create Word document
        doc = Document()
        doc.add_heading(f"{document_type}: {project_name}", 0)
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Check if it's a heading (starts with # or is all caps)
                if para.strip().startswith('#'):
                    heading_text = para.strip().lstrip('#').strip()
                    doc.add_heading(heading_text, level=1)
                elif para.strip().isupper() and len(para.strip()) < 100:
                    doc.add_heading(para.strip(), level=2)
                else:
                    doc.add_paragraph(para.strip())
        
        # Save to BytesIO
        file_data = io.BytesIO()
        doc.save(file_data)
        file_data.seek(0)
        
        filename = f"{safe_project_name}_{safe_doc_type}.docx"
        
    elif file_format == 'xlsx':
        # Create Excel workbook
        wb = Workbook()
        ws = wb.active
        ws.title = document_type[:31]  # Excel sheet name limit
        
        # Add title
        ws['A1'] = f"{document_type}: {project_name}"
        ws['A1'].font = ws['A1'].font.copy(bold=True, size=14)
        
        # Add content
        row = 3
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                ws[f'A{row}'] = line.strip()
                row += 1
        
        # Save to BytesIO
        file_data = io.BytesIO()
        wb.save(file_data)
        file_data.seek(0)
        
        filename = f"{safe_project_name}_{safe_doc_type}.xlsx"
        
    else:  # txt
        file_data = io.BytesIO()
        file_data.write(f"{document_type}: {project_name}\n\n".encode('utf-8'))
        file_data.write(content.encode('utf-8'))
        file_data.seek(0)
        
        filename = f"{safe_project_name}_{safe_doc_type}.txt"
    
    return file_data, filename

def get_mimetype(file_format):
    """Get MIME type for file format"""
    mimetypes = {
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'txt': 'text/plain'
    }
    return mimetypes.get(file_format, 'application/octet-stream')

